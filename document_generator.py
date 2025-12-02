"""
Модуль для генерации маршрутных карт по ГОСТ 3.1118
Форма 4 - первый лист, Форма 3б - последующие листы
"""
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from datetime import datetime

class DocumentGenerator:
    def __init__(self):
        self.doc = None
        # Константы ГОСТ 3.1118
        self.ROW_HEIGHT = Mm(8)  # Высота строки
        self.ROWS_PER_PAGE_FIRST = 15  # Строк на первом листе
        self.ROWS_PER_PAGE_NEXT = 25   # Строк на последующих листах
        
        # Типы строк по ГОСТ
        self.ROW_TYPES = {
            'operation': 'О',  # Операция
            'transition': 'Т',  # Переход
            'equipment': 'М',   # Оборудование
            'material': 'А',    # Материал
            'tool': 'И',        # Инструмент
            'comment': 'К'      # Комментарий
        }
    
    def create_route_card(self, data, output_path, doc_info=None):
        """
        Создание маршрутной карты по ГОСТ 3.1118
        
        Args:
            data: DataFrame с данными элементов и процессов
            output_path: путь для сохранения документа
            doc_info: словарь с информацией о документе (название изделия, обозначение и т.д.)
        """
        self.doc = Document()
        
        # Настройка страницы A4
        section = self.doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        
        # Подготовка данных
        if doc_info is None:
            doc_info = {
                'product_name': 'Печатный узел',
                'designation': '',
                'developer': '',
                'date': datetime.now().strftime('%d.%m.%Y')
            }
        
        # Преобразование данных в строки маршрутной карты
        route_rows = self._prepare_route_data(data)
        
        # Первый лист - Форма 4
        self._add_form_4(route_rows[:self.ROWS_PER_PAGE_FIRST], doc_info)
        
        # Последующие листы - Форма 3б
        remaining_rows = route_rows[self.ROWS_PER_PAGE_FIRST:]
        page_num = 2
        
        while remaining_rows:
            self.doc.add_page_break()
            page_rows = remaining_rows[:self.ROWS_PER_PAGE_NEXT]
            self._add_form_3b(page_rows, doc_info, page_num)
            remaining_rows = remaining_rows[self.ROWS_PER_PAGE_NEXT:]
            page_num += 1
        
        # Сохранение
        self.doc.save(output_path)
    
    def _prepare_route_data(self, data):
        """
        Подготовка данных для маршрутной карты
        Преобразование DataFrame в список строк с типами
        """
        route_rows = []
        
        if data is None or data.empty:
            return route_rows
        
        # Группировка по операциям
        current_operation = None
        row_number = 1
        
        for idx, row in data.iterrows():
            # Получаем значения с проверкой на NaN
            operation = self._get_value(row, ['Operation', 'Процесс', 'operation'])
            designator = self._get_value(row, ['Designator', 'Позиционное обозначение', 'designator'])
            description = self._get_value(row, ['Description', 'Наименование', 'description', 'Comment'])
            quantity = self._get_value(row, ['Quantity', 'Количество', 'quantity'])
            equipment = self._get_value(row, ['Equipment', 'Оборудование', 'equipment'])
            material = self._get_value(row, ['Material', 'Материал', 'material'])
            
            # Если новая операция - добавляем строку операции
            if operation and operation != current_operation and str(operation).lower() != 'nan':
                route_rows.append({
                    'type': 'О',
                    'number': f'О{row_number:02d}',
                    'name': str(operation),
                    'equipment': equipment if equipment else '',
                    'material': material if material else '',
                    'time_prep': '',
                    'time_piece': ''
                })
                current_operation = operation
                row_number += 1
            
            # Формируем наименование элемента
            name_parts = []
            if designator:
                name_parts.append(str(designator))
            if description:
                name_parts.append(str(description))
            
            element_name = ' - '.join(name_parts) if name_parts else 'Элемент'
            
            # Добавляем строку перехода (элемент)
            if designator or description:
                # Обработка количества
                qty_str = '1'
                if quantity and str(quantity).lower() != 'nan':
                    try:
                        # Преобразуем в float, затем в int для удаления дробной части
                        qty_str = str(int(float(quantity)))
                    except (ValueError, TypeError):
                        qty_str = str(quantity)
                
                route_rows.append({
                    'type': 'Т',
                    'number': f'Т{row_number:02d}',
                    'name': element_name,
                    'equipment': equipment if equipment else '',
                    'material': material if material else '',
                    'time_prep': '',
                    'time_piece': qty_str
                })
                row_number += 1
        
        return route_rows
    
    def _get_value(self, row, possible_keys):
        """
        Получение значения из строки по возможным ключам
        Возвращает первое найденное непустое значение или None
        """
        for key in possible_keys:
            if key in row.index:
                value = row[key]
                # Проверка на NaN и пустые значения
                if pd.notna(value) and str(value).strip() and str(value).lower() != 'nan':
                    return str(value).strip()
        return None
    
    def _add_form_4(self, rows, doc_info):
        """
        Добавление первого листа - Форма 4 по ГОСТ 3.1118
        """
        # Заголовок документа
        title = self.doc.add_paragraph()
        title_run = title.add_run('МАРШРУТНАЯ КАРТА')
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информационная рамка (упрощенная версия)
        info_table = self.doc.add_table(rows=4, cols=4)
        info_table.style = 'Table Grid'
        
        # Заполнение информационной рамки
        info_table.rows[0].cells[0].text = 'Наименование изделия'
        info_table.rows[0].cells[1].merge(info_table.rows[0].cells[3])
        info_table.rows[0].cells[1].text = doc_info.get('product_name', '')
        
        info_table.rows[1].cells[0].text = 'Обозначение'
        info_table.rows[1].cells[1].text = doc_info.get('designation', '')
        info_table.rows[1].cells[2].text = 'Лист'
        info_table.rows[1].cells[3].text = '1'
        
        info_table.rows[2].cells[0].text = 'Разработал'
        info_table.rows[2].cells[1].text = doc_info.get('developer', '')
        info_table.rows[2].cells[2].text = 'Дата'
        info_table.rows[2].cells[3].text = doc_info.get('date', '')
        
        info_table.rows[3].cells[0].text = 'Проверил'
        info_table.rows[3].cells[1].text = ''
        info_table.rows[3].cells[2].text = 'Дата'
        info_table.rows[3].cells[3].text = ''
        
        self.doc.add_paragraph()
        
        # Основная таблица маршрутной карты
        self._add_route_table(rows)
    
    def _add_form_3b(self, rows, doc_info, page_num):
        """
        Добавление последующих листов - Форма 3б по ГОСТ 3.1118
        """
        # Упрощенный заголовок для последующих листов
        header_table = self.doc.add_table(rows=2, cols=4)
        header_table.style = 'Table Grid'
        
        header_table.rows[0].cells[0].text = 'Обозначение'
        header_table.rows[0].cells[1].text = doc_info.get('designation', '')
        header_table.rows[0].cells[2].text = 'Лист'
        header_table.rows[0].cells[3].text = str(page_num)
        
        header_table.rows[1].cells[0].merge(header_table.rows[1].cells[3])
        header_table.rows[1].cells[0].text = 'МАРШРУТНАЯ КАРТА (продолжение)'
        
        self.doc.add_paragraph()
        
        # Основная таблица
        self._add_route_table(rows)
    
    def _add_route_table(self, rows):
        """
        Создание основной таблицы маршрутной карты
        Структура по ГОСТ 3.1118
        """
        if not rows:
            self.doc.add_paragraph("Нет данных для отображения")
            return
        
        # Создание таблицы
        # Колонки: Тип | № | Наименование операции/перехода | Оборудование | Материал | Тп.з | Тшт
        table = self.doc.add_table(rows=len(rows) + 1, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки
        headers = ['Тип', '№', 'Наименование операции/перехода', 
                  'Оборудование', 'Материал', 'Тп.з', 'Тшт']
        header_row = table.rows[0]
        
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # Форматирование заголовка
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        
        # Заполнение данных
        for row_idx, row_data in enumerate(rows, start=1):
            table_row = table.rows[row_idx]
            
            # Тип строки
            table_row.cells[0].text = row_data.get('type', '')
            # Номер
            table_row.cells[1].text = row_data.get('number', '')
            # Наименование
            table_row.cells[2].text = row_data.get('name', '')
            # Оборудование
            table_row.cells[3].text = row_data.get('equipment', '')
            # Материал
            table_row.cells[4].text = row_data.get('material', '')
            # Время подготовительно-заключительное
            table_row.cells[5].text = row_data.get('time_prep', '')
            # Время штучное
            table_row.cells[6].text = row_data.get('time_piece', '')
            
            # Форматирование ячеек
            for cell in table_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
        
        # Установка ширины колонок
        widths = [Cm(1), Cm(1.5), Cm(8), Cm(3), Cm(2.5), Cm(1.5), Cm(1.5)]
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = widths[idx]
    
    def set_column_width(self, table):
        """Установка ширины колонок"""
        # Равномерное распределение ширины
        total_width = Cm(17)  # Доступная ширина
        col_width = total_width / len(table.columns)
        
        for row in table.rows:
            for cell in row.cells:
                cell.width = col_width
    
    def add_border(self, cell, **kwargs):
        """Добавление границ к ячейке"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '4')
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), '000000')
            tcBorders.append(edge_element)
        
        tcPr.append(tcBorders)
    
    def add_operation_section(self, operation_name, operation_data):
        """Добавление секции операции"""
        # Заголовок операции
        heading = self.doc.add_heading(operation_name, level=2)
        
        # Таблица операции
        if not operation_data.empty:
            op_table = self.doc.add_table(rows=len(operation_data)+1, 
                                         cols=len(operation_data.columns))
            op_table.style = 'Table Grid'
            
            # Заголовки
            for i, col in enumerate(operation_data.columns):
                op_table.rows[0].cells[i].text = str(col)
            
            # Данные
            for row_idx, (_, row) in enumerate(operation_data.iterrows(), start=1):
                for col_idx, col in enumerate(operation_data.columns):
                    op_table.rows[row_idx].cells[col_idx].text = str(row[col])
    
    def convert_to_pdf(self, docx_path, pdf_path):
        """Конвертация DOCX в PDF"""
        try:
            # Метод 1: Использование docx2pdf (Windows)
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
                return
            except ImportError:
                pass
            
            # Метод 2: Использование reportlab для создания PDF напрямую
            from reportlab.lib.pagesizes import A4
            from reportlab.lib import colors
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # Читаем данные из DOCX
            from docx import Document
            doc = Document(docx_path)
            
            # Создаем PDF
            pdf_doc = SimpleDocTemplate(pdf_path, pagesize=A4)
            elements = []
            styles = getSampleStyleSheet()
            
            # Заголовок
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                alignment=1  # Центрирование
            )
            elements.append(Paragraph("МАРШРУТНАЯ КАРТА", title_style))
            elements.append(Spacer(1, 0.5*cm))
            
            # Извлекаем таблицы из DOCX
            for table in doc.tables:
                data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    data.append(row_data)
                
                if data:
                    # Создаем таблицу в PDF
                    pdf_table = Table(data)
                    pdf_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ]))
                    elements.append(pdf_table)
                    elements.append(Spacer(1, 0.5*cm))
            
            # Сохраняем PDF
            pdf_doc.build(elements)
            
        except Exception as e:
            raise Exception(f"Ошибка конвертации в PDF: {e}")
